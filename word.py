import docx
import requests
import random
import hashlib

# 百度翻译 API 的 URL
url = "https://fanyi-api.baidu.com/api/trans/vip/translate"

# 设置你的 APP ID 和密钥
appid = "your_app_id"
secretKey = "your_secret_key"

def translate_text(text):
    params = {
        "q": text,
        "from": "zh",
        "to": "en",
        "appid": appid,
        "salt": str(random.randint(32768, 65536)),
        "sign": ""
    }
    sign_str = appid + text + params["salt"] + secretKey
    md5 = hashlib.md5()
    md5.update(sign_str.encode('utf-8'))
    params["sign"] = md5.hexdigest()

    try:
        response = requests.post(url, data=params)
        result = response.json()
        if "trans_result" in result:
            return result["trans_result"][0]["dst"]
        else:
            return text
    except requests.RequestException as e:
        print(f"Error during translation: {e}")
        return text

def copy_paragraph_style(paragraph, new_paragraph):
    new_paragraph.style = paragraph.style
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(translate_text(run.text))
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        if run._element.get('w:vertAlign'):
            new_run._element.set('w:vertAlign', run._element.get('w:vertAlign'))

    # Copy paragraph formatting
    new_paragraph.alignment = paragraph.alignment
    new_paragraph.left_indent = paragraph.left_indent
    new_paragraph.right_indent = paragraph.right_indent
    new_paragraph.space_before = paragraph.space_before
    new_paragraph.space_after = paragraph.space_after
    new_paragraph.line_spacing = paragraph.line_spacing

def copy_table_style(table, new_table):
    new_table.style = table.style
    for row in table.rows:
        new_row = new_table.add_row()
        for cell in row.cells:
            new_cell = new_row.cells[len(new_row.cells) - 1]
            new_cell.text = translate_text(cell.text)
            if cell._element.get('w:shd'):
                new_cell._element.set('w:shd', cell._element.get('w:shd'))  # Copy shading (background color)

def copy_header_footer(document, new_document):
    for section in document.sections:
        for header in section.header.part.text_box_content_list:
            new_document.sections[0].header.part.add_textbox(content=header)
        for footer in section.footer.part.text_box_content_list:
            new_document.sections[0].footer.part.add_textbox(content=footer)

def copy_images(document, new_document):
    for shape in document.inline_shapes:
        if shape.type == 3:  # Image type
            image = shape._inline.graphic.graphicData.pic.blip
            image_data = document.part.related_parts[image.embed]
            new_document.add_picture(image_data, shape.width, shape.height)

def translate_docx(input_file, output_file):
    doc = docx.Document(input_file)
    new_doc = docx.Document()

    # Copy header and footer
    copy_header_footer(doc, new_doc)

    # Copy images
    copy_images(doc, new_doc)

    for para in doc.paragraphs:
        new_para = new_doc.add_paragraph()
        copy_paragraph_style(para, new_para)

    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        copy_table_style(table, new_table)

    new_doc.save(output_file)

# 使用示例
translate_docx("path_to_your_input_file.docx", "path_to_your_output_file.docx")